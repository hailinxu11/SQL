'''渲染html并生成新的html'''
from jinja2 import Environment, FileSystemLoader
from jira_html.data import *
from docxtpl import DocxTemplate
from docx.shared import Cm, Inches
import datetime
from .models import *
import os
from datetime import datetime


PATH = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_ENVIRONMENT = Environment(
    autoescape=False,
    loader=FileSystemLoader(os.path.join(PATH, 'static/data')),
    trim_blocks=False)

# 模板路径
path = os.path.join(PATH, 'static/data')
# 获取html文件名
def htmlName():
    for file in os.listdir(path):
        if file.split('.')[-1]== 'html':
            return file

# 获取word文件名
def wordName():
    for file in os.listdir(path):
        if file.split('.')[-1]== 'docx':
            return file

# 渲染html
def render_template(template_filename, context):
    return TEMPLATE_ENVIRONMENT.get_template(template_filename).render(context)


# 创建新的html
def create_index_html(data_id):
    newFileName = datetime.now().strftime("%Y-%m-%d_%H:%M:%S") + '.html'
    fileName = htmlName()
    file_path = PATH + '/static/files/%s' % newFileName
    # 获取功能点信息
    code1 = dataList(data_id)
    # 获取根据模块分组后的功能点
    code2 = codeList(data_id)
    detail = htmlList(data_id)
    context = {
        'code1': code1,
        'code2': code2,
        'Html': detail
    }
    with open(file_path, 'w', encoding='utf-8') as f:
        html = render_template(fileName, context)
        f.write(html)
    # 新生成的文件存入数据库
    File.objects.get_or_create(name=newFileName, path=file_path, file_type='html文件')


# 插入表格
def create_table(document, data):
    table = document.add_table(rows=1, cols=4, style="Table Grid")
    table.autofit = True
    # 设置功能说明列的列宽
    table.columns[2].width = Cm(6)
    # 设置备注列宽
    table.columns[3].width = Cm(2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '一级功能'
    hdr_cells[1].text = '二级功能'
    hdr_cells[2].text = '功能说明'
    hdr_cells[3].text = '备注'
    last_epic_name = None
    last_epic_index = 1
    for d in data:
        if last_epic_name is not None and last_epic_name != d[0]:
            row_count = len(table.rows)
            # 合并单元格
            merages = table.cell(last_epic_index, 0).merge(table.cell(row_count - 1, 0))
            # 往合并的单元格中填充数据
            merages.text = last_epic_name
            # 设置合并单元格的段落缩进
            merages.paragraphs[0].paragraph_format.left_indent = Inches(0.2)
            last_epic_index = row_count
        # 赋值变量为一级功能名称
        last_epic_name = d[0]
        row_cells = table.add_row().cells
        row_cells[0].text = d[0]
        row_cells[1].text = d[1]
        row_cells[2].text = d[2].strip() # 去空格
        # 设置二级功能的堕落缩进
        table.cell(len(table.rows)-1, 1).paragraphs[0].paragraph_format.left_indent = Inches(0.2)
    if last_epic_index != len(table.rows):
        # 合并单元格
        merages = table.cell(last_epic_index, 0).merge(table.cell(len(table.rows) - 1, 0))
        # 往合并的单元格中填充数据
        merages.text = last_epic_name
        # 设置合并单元格的段落缩进
        merages.paragraphs[0].paragraph_format.left_indent = Inches(0.2)
        # table.cell(last_epic_index, 0).merge(table.cell(len(table.rows) - 1, 0)).text = last_epic_name


# 写入数据至word
def write_to_word(data_id):
    newFileName = datetime.now().strftime("%Y-%m-%d_%H:%M:%S") + '.docx'
    fileName = wordName()
    filePath = os.path.join(path, fileName)
    # 定义当前时间
    codes = wordList(data_id)
    doc = DocxTemplate(filePath)
    i = 1
    for key, val in codes.items():
        doc.add_paragraph(str(i) + '.' + key)
        doc.add_paragraph('')
        create_table(doc, val)
        doc.add_paragraph('')
        i += 1

    doc.add_paragraph('''

    （以下无正文）





甲方：                               乙方：

法定代表人：（签字）                 法定代表人：（签字）

委托代理人：                         委托代理人：

开户银行；                           开户银行：

银行账号：                           银行账号：

签订日期：                           签订日期：
        ''')
    newFilePath = PATH + '/static/files/%s' % newFileName
    doc.save(newFilePath)
    File.objects.get_or_create(name=newFileName, path=newFilePath, file_type="word文件")
